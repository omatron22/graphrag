#!/usr/bin/env python3
"""
Fixed test script for the Knowledge Graph-Based Business Consulting System.
This script corrects the model name checking logic and creates a proper PDF test document.
"""

import os
import sys
import logging
import json
import requests
import importlib
import time
import csv
from datetime import datetime

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("system_test.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("system_test")

def check_directories():
    """Check if all required directories exist"""
    required_dirs = [
        "neo4j/data",
        "neo4j/logs", 
        "neo4j/import",
        "neo4j/plugins",
        "postgres-data",
        "data/uploads",
        "data/parsed",
        "data/knowledge_base",
        "knowledge_graph/exports"
    ]
    
    missing_dirs = []
    for directory in required_dirs:
        if not os.path.exists(directory):
            missing_dirs.append(directory)
            
    if missing_dirs:
        logger.error(f"Missing directories: {', '.join(missing_dirs)}")
        print(f"❌ Directory check failed. Missing: {', '.join(missing_dirs)}")
        return False
    else:
        logger.info("All required directories exist")
        print("✅ Directory check passed")
        return True

def check_neo4j_connection():
    """Test connection to Neo4j database"""
    try:
        # Import Neo4j modules
        from neo4j import GraphDatabase
        from knowledge_graph.neo4j_manager import Neo4jManager
        
        # Get connection details from config
        import config
        
        # Create Neo4jManager instance
        neo4j_manager = Neo4jManager(
            uri=config.NEO4J_URI,
            user=config.NEO4J_USER,
            password=config.NEO4J_PASSWORD
        )
        
        # Test connection
        connected = neo4j_manager.connect()
        if connected:
            logger.info("Successfully connected to Neo4j")
            print("✅ Neo4j connection test passed")
            
            # Test basic query
            result = neo4j_manager.execute_query("RETURN 'Test' AS test")
            if result and result[0].get('test') == 'Test':
                logger.info("Neo4j query test passed")
                print("✅ Neo4j query test passed")
            else:
                logger.error("Neo4j query test failed")
                print("❌ Neo4j query test failed")
                connected = False
                
            neo4j_manager.close()
            return connected
        else:
            logger.error("Failed to connect to Neo4j")
            print("❌ Neo4j connection test failed")
            return False
            
    except Exception as e:
        logger.error(f"Neo4j connection test error: {str(e)}")
        print(f"❌ Neo4j connection test error: {str(e)}")
        return False

def check_postgres_connection():
    """Test connection to PostgreSQL database"""
    try:
        # Import PostgreSQL modules
        import psycopg2
        
        # Get connection details from config
        import config
        
        # Test connection
        conn = psycopg2.connect(config.POSTGRES_URI)
        cursor = conn.cursor()
        cursor.execute("SELECT 1")
        result = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if result and result[0] == 1:
            logger.info("Successfully connected to PostgreSQL")
            print("✅ PostgreSQL connection test passed")
            return True
        else:
            logger.error("PostgreSQL query test failed")
            print("❌ PostgreSQL query test failed")
            return False
            
    except Exception as e:
        logger.error(f"PostgreSQL connection test error: {str(e)}")
        print(f"❌ PostgreSQL connection test error: {str(e)}")
        return False

def check_ollama_service():
    """Test connection to Ollama service and check if required models are available"""
    try:
        # Get Ollama endpoint from config
        import config
        
        # Get required model names directly from config and clean them
        required_models = []
        for model_type in ['reasoning', 'vision', 'lightweight']:
            model_name = config.MODELS[model_type]['name']
            # Remove any trailing colons that shouldn't be there
            if model_name.endswith(':'):
                logger.warning(f"Model name '{model_name}' has a trailing colon, which is incorrect")
                model_name = model_name.rstrip(':')
            required_models.append(model_name)
        
        # Log what we're looking for
        logger.info(f"Required models (cleaned): {required_models}")
        
        # Check if Ollama service is running
        try:
            response = requests.get("http://localhost:11434/api/tags")
            if response.status_code == 200:
                available_models = response.json().get("models", [])
                
                # Extract just the model names and model:tag formats
                available_model_names = set()
                for model in available_models:
                    name = model.get("name")
                    if name:
                        available_model_names.add(name)
                        # Also add the base name without tag if it has one
                        if ":" in name:
                            base_name = name.split(":")[0]
                            available_model_names.add(base_name)
                
                logger.info(f"Ollama service is running. Available models: {list(available_model_names)}")
                print("✅ Ollama service is running")
                
                # Check required models with more flexible matching
                missing_models = []
                for model in required_models:
                    base_model = model.split(":")[0] if ":" in model else model
                    
                    # Check if the model name or a tagged version is available
                    if (model not in available_model_names and 
                        base_model not in available_model_names and 
                        not any(m.startswith(f"{base_model}:") for m in available_model_names)):
                        missing_models.append(model)
                
                if missing_models:
                    logger.warning(f"Missing required models: {missing_models}")
                    print(f"⚠️ Missing required models: {', '.join(missing_models)}")
                    print(f"   Please install them with: ollama pull {' '.join(missing_models)}")
                    return False
                else:
                    logger.info("All required models are available")
                    print("✅ All required Ollama models are available")
                    return True
            else:
                logger.error(f"Ollama service returned status code: {response.status_code}")
                print(f"❌ Ollama service returned unexpected status: {response.status_code}")
                return False
                
        except requests.exceptions.ConnectionError:
            logger.error("Could not connect to Ollama service")
            print("❌ Ollama service connection failed - is Ollama running?")
            print("   Start Ollama with: ollama serve")
            return False
            
    except Exception as e:
        logger.error(f"Ollama service test error: {str(e)}")
        print(f"❌ Ollama service test error: {str(e)}")
        return False

def create_test_document():
    """Create a test document with rich business information in a supported format"""
    test_dir = "data/uploads"
    os.makedirs(test_dir, exist_ok=True)
    
    # Import config to check supported formats
    import config
    supported_formats = list(config.EXTRACTORS.keys())
    logger.info(f"Supported document formats: {supported_formats}")
    
    # Try creating a PDF with reportlab
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        test_file = os.path.join(test_dir, "test_business_report.pdf")
        
        # Create PDF document
        c = canvas.Canvas(test_file, pagesize=letter)
        c.setFont("Helvetica", 12)
        
        # Add title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(72, 750, "Business Performance Report - Q1 2025")
        c.setFont("Helvetica", 12)
        
        # Add content sections
        y_position = 720
        
        # Executive Summary
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, y_position, "1. Executive Summary:")
        c.setFont("Helvetica", 12)
        y_position -= 20
        c.drawString(90, y_position, "TechCorp Inc. has shown strong financial performance in Q1 2025.")
        y_position -= 15
        c.drawString(90, y_position, "Revenue increased by 15% compared to the previous quarter,")
        y_position -= 15
        c.drawString(90, y_position, "reaching $120 million. Net profit margin improved to 18% from 15%.")
        
        # Financial Highlights
        y_position -= 30
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, y_position, "2. Financial Highlights:")
        c.setFont("Helvetica", 12)
        y_position -= 20
        c.drawString(90, y_position, "- Revenue: $120 million (↑15% QoQ)")
        y_position -= 15
        c.drawString(90, y_position, "- Operating Expenses: $75 million (↑5% QoQ)")
        y_position -= 15
        c.drawString(90, y_position, "- EBITDA: $42 million (↑28% QoQ)")
        y_position -= 15
        c.drawString(90, y_position, "- Net Profit: $21.6 million (↑38% QoQ)")
        y_position -= 15
        c.drawString(90, y_position, "- Cash Reserves: $85 million (↑10% QoQ)")
        
        # Market Performance
        y_position -= 30
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, y_position, "3. Market Performance:")
        c.setFont("Helvetica", 12)
        y_position -= 20
        c.drawString(90, y_position, "TechCorp Inc. has gained 2.3% market share in the enterprise")
        y_position -= 15
        c.drawString(90, y_position, "software segment, surpassing CompetitorX. The company now ranks")
        y_position -= 15
        c.drawString(90, y_position, "second in the industry behind MarketLeader Corp.")
        
        # Product Performance
        y_position -= 30
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, y_position, "4. Product Performance:")
        c.setFont("Helvetica", 12)
        y_position -= 20
        c.drawString(90, y_position, "- CloudSuite Pro: Revenue of $65 million (↑22% QoQ)")
        y_position -= 15
        c.drawString(90, y_position, "- SecureNet Enterprise: Revenue of $32 million (↑8% QoQ)")
        y_position -= 15
        c.drawString(90, y_position, "- DataAnalyzer: Revenue of $23 million (↑12% QoQ)")
        
        # Strategic Initiatives
        y_position -= 30
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, y_position, "5. Strategic Initiatives:")
        c.setFont("Helvetica", 12)
        y_position -= 20
        c.drawString(90, y_position, "TechCorp Inc. has signed a partnership agreement with GlobalNet Inc.")
        y_position -= 15
        c.drawString(90, y_position, "to expand into Asian markets. The company acquired SmallTech Solutions")
        y_position -= 15
        c.drawString(90, y_position, "for $35 million to strengthen its AI capabilities.")
        
        # Save the PDF
        c.save()
        
        logger.info(f"Created rich test document: {test_file}")
        print(f"✅ Created rich test document: {test_file}")
        return test_file
        
    except ImportError:
        logger.warning("ReportLab not installed. Trying alternative methods...")
        print("⚠️ ReportLab not installed. Trying alternative methods...")
    except Exception as e:
        logger.error(f"Error creating PDF with ReportLab: {e}")
    
    # Try creating a DOCX if supported
    if 'docx' in supported_formats:
        try:
            from docx import Document
            
            test_file = os.path.join(test_dir, "test_business_report.docx")
            
            doc = Document()
            doc.add_heading('Business Performance Report - Q1 2025', 0)
            
            doc.add_heading('1. Executive Summary:', level=1)
            doc.add_paragraph('TechCorp Inc. has shown strong financial performance in Q1 2025. '
                            'Revenue increased by 15% compared to the previous quarter, '
                            'reaching $120 million. Net profit margin improved to 18% from 15% in Q4 2024.')
            
            doc.add_heading('2. Financial Highlights:', level=1)
            p = doc.add_paragraph()
            p.add_run('- Revenue: $120 million (↑15% QoQ)\n')
            p.add_run('- Operating Expenses: $75 million (↑5% QoQ)\n')
            p.add_run('- EBITDA: $42 million (↑28% QoQ)\n')
            p.add_run('- Net Profit: $21.6 million (↑38% QoQ)\n')
            p.add_run('- Cash Reserves: $85 million (↑10% QoQ)')
            
            doc.add_heading('3. Market Performance:', level=1)
            doc.add_paragraph('TechCorp Inc. has gained 2.3% market share in the enterprise software segment, '
                            'surpassing CompetitorX. The company now ranks second in the industry behind '
                            'MarketLeader Corp.')
            
            doc.add_heading('4. Strategic Initiatives:', level=1)
            doc.add_paragraph('TechCorp Inc. has signed a partnership agreement with GlobalNet Inc. '
                            'to expand into Asian markets. The company acquired SmallTech Solutions '
                            'for $35 million to strengthen its AI capabilities.')
            
            # Save the document
            doc.save(test_file)
            
            logger.info(f"Created DOCX document: {test_file}")
            print(f"✅ Created DOCX document: {test_file}")
            return test_file
            
        except ImportError:
            logger.warning("python-docx not installed.")
            print("⚠️ python-docx not installed.")
        except Exception as e:
            logger.error(f"Error creating DOCX document: {e}")
    
    # Create a simple text-based solution for each supported format
    for format in supported_formats:
        try:
            # Simple content for the document
            content = """Business Performance Report - Q1 2025

1. Executive Summary:
   TechCorp Inc. has shown strong financial performance in Q1 2025. Revenue increased by 15% compared to the previous quarter, reaching $120 million. Net profit margin improved to 18% from 15% in Q4 2024.

2. Financial Highlights:
   - Revenue: $120 million (↑15% QoQ)
   - Operating Expenses: $75 million (↑5% QoQ)
   - EBITDA: $42 million (↑28% QoQ)
   - Net Profit: $21.6 million (↑38% QoQ)
   - Cash Reserves: $85 million (↑10% QoQ)

3. Strategic Initiatives:
   TechCorp Inc. has signed a partnership agreement with GlobalNet Inc. to expand into Asian markets. The company acquired SmallTech Solutions for $35 million to strengthen its AI capabilities.
"""
            
            # Create file based on format
            if format == 'pdf' and 'pdf' in supported_formats:
                # Try a very minimal PDF creation approach
                try:
                    # Use a very basic approach to create a PDF
                    test_file = os.path.join(test_dir, "test_business_report.pdf")
                    
                    # Use a very simple way to create PDF without third-party libraries
                    with open(test_file, 'wb') as f:
                        f.write(b'%PDF-1.7\n\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n\n2 0 obj\n<<\n/Type /Pages\n/Kids [3 0 R]\n/Count 1\n>>\nendobj\n\n3 0 obj\n<<\n/Type /Page\n/Parent 2 0 R\n/Resources <<\n/Font <<\n/F1 4 0 R\n>>\n>>\n/MediaBox [0 0 612 792]\n/Contents 5 0 R\n>>\nendobj\n\n4 0 obj\n<<\n/Type /Font\n/Subtype /Type1\n/BaseFont /Helvetica\n>>\nendobj\n\n5 0 obj\n<<\n/Length 68\n>>\nstream\nBT\n/F1 12 Tf\n100 700 Td\n(Business Performance Report - Q1 2025) Tj\nET\nendstream\nendobj\n\nxref\n0 6\n0000000000 65535 f\n0000000010 00000 n\n0000000059 00000 n\n0000000116 00000 n\n0000000251 00000 n\n0000000318 00000 n\n\ntrailer\n<<\n/Size 6\n/Root 1 0 R\n>>\n\nstartxref\n436\n%%EOF')
                    
                    logger.info(f"Created minimal PDF document: {test_file}")
                    print(f"✅ Created minimal PDF document: {test_file}")
                    return test_file
                    
                except Exception as e:
                    logger.error(f"Error creating minimal PDF: {e}")
            
            elif format in ['xlsx', 'csv']:
                # Create a very simple CSV
                test_file = os.path.join(test_dir, f"test_business_report.{format}")
                
                with open(test_file, 'w', newline='') as f:
                    writer = csv.writer(f)
                    writer.writerow(['Metric', 'Value', 'Change'])
                    writer.writerow(['Revenue', '$120 million', '+15%'])
                    writer.writerow(['Operating Expenses', '$75 million', '+5%'])
                    writer.writerow(['EBITDA', '$42 million', '+28%'])
                    writer.writerow(['Net Profit', '$21.6 million', '+38%'])
                    writer.writerow(['Cash Reserves', '$85 million', '+10%'])
                    
                logger.info(f"Created {format.upper()} document: {test_file}")
                print(f"✅ Created {format.upper()} document: {test_file}")
                return test_file
        
        except Exception as e:
            logger.error(f"Error creating {format} document: {e}")
    
    # If we get here, we couldn't create any supported document
    logger.error("Failed to create any test document in a supported format")
    print("❌ Failed to create a test document in a supported format")
    print(f"   Supported formats: {', '.join(supported_formats)}")
    return None

def test_document_extraction():
    """Test document extraction with a better test document"""
    try:
        # Create a test document
        test_path = create_test_document()
        
        if not test_path:
            logger.error("Could not create a test document in a supported format")
            print("❌ Document creation failed - could not create a document in a supported format")
            return False
        
        # Ensure the test file has a supported extension by checking config
        import config
        file_ext = os.path.splitext(test_path)[1].lower()[1:]  # Remove the dot
        
        if file_ext not in config.EXTRACTORS:
            logger.error(f"Created test file has unsupported extension: {file_ext}")
            print(f"❌ Test file has unsupported extension: {file_ext}")
            print(f"   Supported extensions: {', '.join(config.EXTRACTORS.keys())}")
            return False
        
        # Import orchestrator
        from orchestrator import Orchestrator
        
        # Initialize orchestrator
        orchestrator = Orchestrator()
        
        # Test document processing
        logger.info(f"Testing document processing with: {test_path}")
        print(f"Testing document processing (this may take a minute)...")
        
        try:
            result = orchestrator.process_document(test_path)
            
            if result.get("status") == "complete":
                logger.info(f"Document processing successful: {result}")
                print(f"✅ Document processing test passed")
                
                # Check if triplets were extracted
                if result.get("primary_entities"):
                    print(f"   Primary entities found: {', '.join(result.get('primary_entities'))}")
                else:
                    print(f"   No primary entities found - knowledge graph may be empty")
                
                return True
            else:
                logger.error(f"Document processing failed: {result}")
                print(f"❌ Document processing test failed")
                if "error" in result:
                    print(f"   Error: {result['error']}")
                return False
                
        except Exception as e:
            logger.error(f"Document processing test error: {str(e)}")
            print(f"❌ Document processing test error: {str(e)}")
            return False
        finally:
            # Clean up
            orchestrator.cleanup()
            
    except Exception as e:
        logger.error(f"Document extraction test setup error: {str(e)}")
        print(f"❌ Document extraction test setup error: {str(e)}")
        return False

def fix_config_model_names():
    """Fix any issues with model names in config.py"""
    try:
        import config
        
        # Check for and fix any trailing colons in model names
        fixed_config = False
        
        for model_type, model_config in config.MODELS.items():
            model_name = model_config['name']
            if model_name.endswith(':'):
                logger.warning(f"Found model name with trailing colon: {model_name}")
                # Fix the issue in memory
                config.MODELS[model_type]['name'] = model_name.rstrip(':')
                fixed_config = True
                logger.info(f"Fixed model name in memory: {model_name} -> {config.MODELS[model_type]['name']}")
                
        if fixed_config:
            print("⚠️ Found and fixed model name issues in memory")
            print("   You should update your config.py file to remove trailing colons from model names")
        
        return fixed_config
        
    except Exception as e:
        logger.error(f"Error checking/fixing config model names: {e}")
        return False

def main():
    """Run all tests and print summary"""
    print("\n===== Knowledge Graph Business Consulting System Test =====\n")
    
    # Track test results
    results = {}
    
    # Check Python version
    python_version = sys.version.split()[0]
    logger.info(f"Python version: {python_version}")
    print(f"Python version: {python_version}")
    
    # Import config to check model names
    try:
        import config
        logger.info(f"Lightweight model in config: {config.MODELS['lightweight']['name']}")
        logger.info(f"Vision model in config: {config.MODELS['vision']['name']}")
        logger.info(f"Reasoning model in config: {config.MODELS['reasoning']['name']}")
        
        # Fix any issues with model names
        fix_config_model_names()
        
    except Exception as e:
        logger.error(f"Error importing config: {e}")
    
    # Run tests
    print("\n----- Running Tests -----\n")
    
    results["directories"] = check_directories()
    results["neo4j"] = check_neo4j_connection()
    results["postgres"] = check_postgres_connection()
    results["ollama"] = check_ollama_service()
    results["extraction"] = test_document_extraction()
    
    # Print summary
    print("\n----- Test Summary -----\n")
    
    all_passed = all(results.values())
    
    for test, passed in results.items():
        status = "✅ PASSED" if passed else "❌ FAILED"
        print(f"{test.ljust(15)}: {status}")
    
    print("\n----- Overall Result -----\n")
    
    if all_passed:
        print("✅ All tests PASSED - System is correctly configured")
        print("\nYou can now run the system with: python run_system.py")
    else:
        print("❌ Some tests FAILED - Please fix the issues noted above")
        print("\nCheck system_test.log for more detailed error information")
    
    print("\nFor more help, refer to the README.md file or check the project documentation")
    print("\n===============================================================\n")
    
    return 0 if all_passed else 1

if __name__ == "__main__":
    sys.exit(main())