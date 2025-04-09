"""
DOCX document extractor for the knowledge graph-based business consulting system.
Extracts text, tables, and structured information from Word documents.
"""

import os
import logging
import json
from datetime import datetime
import docx
from extractors.extractor_utils import (
    ensure_output_dir, generate_output_filename, save_extraction_result,
    create_extraction_result_template, detect_entities, COMMON_ENTITY_PATTERNS,
    extract_document_date, clean_text
)

# Set up logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def extract(file_path, output_dir=None):
    """
    Extract content from a DOCX document
    
    Args:
        file_path (str): Path to the DOCX file
        output_dir (str, optional): Directory to save parsed data
        
    Returns:
        dict: Extracted content with metadata
    """
    logger.info(f"Extracting content from DOCX: {file_path}")
    
    # Validate file exists and is a DOCX
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    if not file_path.lower().endswith('.docx'):
        raise ValueError(f"File is not a DOCX: {file_path}")
    
    # Ensure output directory exists
    output_dir = ensure_output_dir(output_dir)
    
    # Generate output filename
    name_without_ext, output_filename = generate_output_filename(file_path)
    output_path = os.path.join(output_dir, output_filename)
    
    # Create extraction result template
    result = create_extraction_result_template(file_path, "docx_extractor")
    
    try:
        # Open the DOCX document
        doc = docx.Document(file_path)
        
        # Extract document metadata
        result["document_metadata"] = {
            "title": doc.core_properties.title or name_without_ext,
            "author": doc.core_properties.author or "",
            "subject": doc.core_properties.subject or "",
            "keywords": doc.core_properties.keywords or "",
            "created": doc.core_properties.created.isoformat() if doc.core_properties.created else "",
            "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else "",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
        
        # Extract text content
        full_text = ""
        paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            text = clean_text(para.text)
            if text.strip():  # Skip empty paragraphs
                paragraphs.append({
                    "paragraph_number": i + 1,
                    "text": text,
                    "is_heading": para.style.name.startswith('Heading')
                })
                full_text += text + "\n"
        
        # Process tables
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [clean_text(cell.text) for cell in row.cells]
                table_data.append(row_data)
                
            if table_data:  # Skip empty tables
                tables.append({
                    "table_number": i + 1,
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0,
                    "data": table_data
                })
        
        # Detect entities in the full text
        entities = detect_entities(full_text, COMMON_ENTITY_PATTERNS)
        
        # Try to extract document date
        document_date = extract_document_date(full_text)
        if document_date:
            result["document_metadata"]["document_date"] = document_date
        
        # Add all extracted content to the result
        result["content"]["text"] = full_text
        result["content"]["paragraphs"] = paragraphs
        result["content"]["tables"] = tables
        result["content"]["entities"] = entities
        
        # Save extraction result
        save_extraction_result(result, output_path)
        
        logger.info(f"Extraction completed successfully. Output saved to: {output_path}")
        return result
    
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise

# For testing
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python docx_extractor.py <docx_file_path>")
        sys.exit(1)
        
    docx_path = sys.argv[1]
    extract(docx_path)